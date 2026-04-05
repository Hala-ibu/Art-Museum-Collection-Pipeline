import csv
import xml.etree.ElementTree as ET
import os
import json
import sys
import pdfplumber
import re
from docx import Document
from openpyxl import load_workbook
from datetime import datetime
from storage.mongo import save_to_mongo 
import chardet

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'src')))



def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    confidence = result.get("confidence")

    print(f"Detected encoding: {encoding} (confidence: {confidence})")

    text = raw.decode(encoding, errors="replace")

    return text

def extract_data_from_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb["Museum Inventory"] 

    art_pieces = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        # Skip row if the first cell is empty
        if row[0] is None: 
            continue 
        
        # Helper function to get data safely
        def get_val(index):
            return row[index] if len(row) > index else None
        
        piece = {
            "id": get_val(0),
            "title": get_val(1),
            "artist": get_val(2),
            "period": get_val(3),
            "type": get_val(4),
            "acquisition_cost": get_val(5),
            "estimated_value": get_val(6),
            "rating": get_val(7),          
            "country_of_origin": get_val(8)
        }
        art_pieces.append(piece)

    return art_pieces
def extract_summary_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Summary"]

    summary = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is not None:
            metric = row[0]
            value = row[1]
            summary[metric] = value

    return summary

def extract_text_from_pdf(pdf_path):
    pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(normalize_text(text))

    return "\n\n".join(pages)

def normalize_text(text):
    if not text:
        return ""

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()

def extract_text_from_two_column_pdf(pdf_path, gap=10):
    pages_text = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            mid_x = page.width / 2

            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))

            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")

            combined = "\n\n".join(part for part in [left_text, right_text] if part)
            if combined:
                pages_text.append(combined)

    return "\n\n".join(pages_text)

def extract_text_from_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

def extract_text_from_two_column_word(docx_path):
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

if __name__ == "__main__":

    pdf_path = "../../data/raw/pdf/art_tables_documentpdf.pdf"
    docx_path="../../data/raw/word/art_tables_document.docx"
    excel_path="../../data/raw/excel/Book1.xlsx"


    pdf_art = extract_text_from_pdf(pdf_path)
    print(pdf_art)
    word_art=extract_text_from_word(docx_path)
    print(word_art)
    excel_art=extract_data_from_excel(excel_path)
    print(excel_art)
    sum_art=extract_summary_from_excel(excel_path)
    print(sum_art)


